from __future__ import annotations

import hashlib
import json
import re
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document


SUPPORTED_SUFFIXES = {".docx", ".pages"}
EXCLUDED_DIRECTORIES = {"output", "state", "work"}
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
UNSUPPORTED_PARTS = {
    "word/comments.xml": "Документ содержит комментарии.",
    "word/people.xml": "Документ содержит данные совместного редактирования.",
}
UNSUPPORTED_XML = {
    "ins": "Документ содержит отслеживаемые вставки.",
    "del": "Документ содержит отслеживаемые удаления.",
    "txbxContent": "Документ содержит связанные текстовые блоки.",
}


def natural_key(path: Path) -> tuple:
    return tuple(
        int(part) if part.isdigit() else part.casefold()
        for part in re.split(r"(\d+)", path.name)
    )


def discover_chapters(project_dir: Path) -> list[Path]:
    project_dir = project_dir.resolve()
    input_dir = project_dir / "input"
    root_files = [
        path
        for path in project_dir.iterdir()
        if path.is_file() and path.suffix.casefold() in SUPPORTED_SUFFIXES
    ]
    input_files = [] if not input_dir.is_dir() else [
        path
        for path in input_dir.iterdir()
        if path.is_file() and path.suffix.casefold() in SUPPORTED_SUFFIXES
    ]
    if root_files and input_files:
        raise ValueError(
            "Поддерживаемые документы найдены одновременно в input/ и корне проекта."
        )
    chapters = input_files or root_files
    folded = [path.stem.casefold() for path in chapters]
    if len(folded) != len(set(folded)):
        raise ValueError("Имена глав неоднозначны с учётом регистра.")
    keys = [natural_key(path) for path in chapters]
    if len(keys) != len(set(keys)):
        raise ValueError("Порядок глав неоднозначен из-за совпадающих числовых частей.")
    if not chapters:
        raise ValueError("Поддерживаемые главы .docx или .pages не найдены.")
    return sorted(chapters, key=natural_key)


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _relative_path(project_dir: Path, path: Path) -> str:
    project_dir = project_dir.resolve()
    resolved = path.resolve()
    try:
        return resolved.relative_to(project_dir).as_posix()
    except ValueError as error:
        raise ValueError("Исходная глава должна находиться внутри проекта.") from error


def _write_manifest(path: Path, manifest: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    temporary.replace(path)


def build_manifest(project_dir: Path, chapters: list[Path]) -> dict:
    project_dir = project_dir.resolve()
    ordered = sorted((path.resolve() for path in chapters), key=natural_key)
    entries = []
    for number, chapter in enumerate(ordered, start=1):
        relative = _relative_path(project_dir, chapter)
        entries.append(
            {
                "номер": number,
                "имя": chapter.name,
                "путь": relative,
                "sha256": file_sha256(chapter),
            }
        )

    parents = {Path(entry["путь"]).parent.as_posix() for entry in entries}
    source = next(iter(parents), ".") if len(parents) == 1 else "смешанный"
    manifest = {"версия": 1, "источник": source, "главы": entries}
    _write_manifest(project_dir / "work" / "manifest.json", manifest)
    return manifest


def _manifest_paths(manifest: dict) -> list[str]:
    chapters = manifest.get("главы")
    if not isinstance(chapters, list):
        raise ValueError("Манифест не содержит корректный список глав.")
    paths = []
    for chapter in chapters:
        if not isinstance(chapter, dict) or not isinstance(chapter.get("путь"), str):
            raise ValueError("Манифест содержит главу без корректного пути.")
        paths.append(Path(chapter["путь"]).as_posix())
    return paths


def verify_manifest(project_dir: Path, manifest: dict) -> list[str]:
    project_dir = project_dir.resolve()
    try:
        manifest_paths = _manifest_paths(manifest)
    except ValueError as error:
        return [str(error)]

    try:
        current = discover_chapters(project_dir)
    except ValueError as error:
        if "не найдены" in str(error):
            current = []
        else:
            return [str(error)]

    current_paths = {_relative_path(project_dir, path): path for path in current}
    errors = []
    for chapter in manifest["главы"]:
        relative = Path(chapter["путь"]).as_posix()
        path = current_paths.get(relative)
        if path is None:
            errors.append(f"Глава «{chapter.get('имя', relative)}» удалена.")
            continue
        expected = chapter.get("sha256")
        if not isinstance(expected, str) or file_sha256(path) != expected:
            errors.append(f"Глава «{path.name}» изменена.")

    current_order = [_relative_path(project_dir, path) for path in current]
    expected_order = manifest_paths
    new_paths = [path for path in current_order if path not in expected_order]
    if expected_order:
        last_key = natural_key(Path(expected_order[-1]))
        if any(natural_key(Path(path)) <= last_key for path in new_paths):
            errors.append(
                "Новые главы должны добавляться только после последней зафиксированной главы."
            )
    return errors


def _output_suffix(chapter: Path, output_format: str) -> str | None:
    selected = output_format.strip().casefold()
    if selected in {"как в оригинале", "оригинал", "original"}:
        return chapter.suffix.casefold()
    if selected in {"docx", ".docx"}:
        return ".docx"
    if selected in {"pages", ".pages"}:
        return ".pages"
    return None


def _confirmed_outputs(
    project_dir: Path, chapters: list[Path], output_format: str
) -> tuple[set[str], str | None, str | None]:
    progress_path = project_dir / "progress.json"
    try:
        progress = json.loads(progress_path.read_text(encoding="utf-8"))
    except (OSError, UnicodeDecodeError, json.JSONDecodeError):
        return set(), None, None
    if not isinstance(progress, dict):
        return set(), None, None

    last_ready = progress.get("последняя_готовая_глава")
    if last_ready is None:
        return set(), None, None
    ordered = sorted(chapters, key=natural_key)
    try:
        last_index = next(
            index for index, chapter in enumerate(ordered)
            if chapter.name == last_ready
        )
    except StopIteration:
        return set(), None, None
    confirmed = set()
    last_name = None
    for chapter in ordered[: last_index + 1]:
        suffix = _output_suffix(chapter, output_format)
        if suffix is None:
            return set(), None, None
        name = f"{chapter.stem}{suffix}"
        confirmed.add(name.casefold())
        last_name = name
    result_hash = progress.get("sha256_результата")
    if result_hash is not None and not isinstance(result_hash, str):
        return set(), None, None
    return confirmed, last_name.casefold() if last_name else None, result_hash


def check_output_conflicts(
    project_dir: Path, chapters: list[Path], output_format: str
) -> list[str]:
    project_dir = project_dir.resolve()
    output_dir = project_dir / "output"
    names = {}
    errors = []
    confirmed, last_confirmed, result_hash = _confirmed_outputs(
        project_dir, chapters, output_format
    )
    for chapter in sorted(chapters, key=natural_key):
        suffix = _output_suffix(chapter, output_format)
        if suffix is None:
            errors.append(f"Неподдерживаемый выходной формат: {output_format}.")
            break
        name = f"{chapter.stem}{suffix}"
        folded = name.casefold()
        previous = names.get(folded)
        if previous is not None:
            errors.append(
                f"Главы «{previous.name}» и «{chapter.name}» создают одно имя результата «{name}»."
            )
        else:
            names[folded] = chapter

        if output_dir.exists():
            existing = next(
                (path for path in output_dir.iterdir() if path.name.casefold() == folded),
                None,
            )
            if existing is not None and existing.name.casefold() not in confirmed:
                errors.append(
                    f"Выходной файл «{existing.name}» уже существует и не подтверждён текущей контрольной точкой."
                )
            elif (
                existing is not None
                and existing.name.casefold() == last_confirmed
                and result_hash is not None
                and file_sha256(existing) != result_hash
            ):
                errors.append(
                    f"Выходной файл «{existing.name}» не совпадает с контрольной суммой контрольной точки."
                )
    return errors


def _block_type(paragraph) -> str:
    text = paragraph.text.strip()
    style = (paragraph.style.name or "") if paragraph.style else ""
    normalized_style = style.casefold()
    if normalized_style.startswith(("heading", "заголовок")):
        return "заголовок"
    if text in {"***", "* * *", "— — —"} or "scene" in normalized_style:
        return "разрыв_сцены"
    return "абзац"


def _run_footnote_id(run) -> str | None:
    element = getattr(run, "_r", run)
    for reference in element.iter(f"{{{W}}}footnoteReference"):
        return reference.get(f"{{{W}}}id")
    return None


def _run_fragments(run) -> list[dict]:
    fragment = {
        "текст": run.text,
        "курсив": bool(run.italic),
        "полужирный": bool(run.bold),
        "сноска": _run_footnote_id(run),
    }
    if fragment["сноска"] is None or not fragment["текст"]:
        return [fragment]
    return [{**fragment, "сноска": None}, {**fragment, "текст": ""}]


def _footnote_fragments(paragraph) -> tuple[str, list[dict]]:
    style = "Normal"
    style_element = paragraph.find(f"./{{{W}}}pPr/{{{W}}}pStyle")
    if style_element is not None:
        style = style_element.get(f"{{{W}}}val", style)
    fragments = []
    for run in paragraph.findall(f"./{{{W}}}r"):
        properties = run.find(f"./{{{W}}}rPr")

        def has_property(name: str) -> bool:
            if properties is None:
                return False
            element = properties.find(f"./{{{W}}}{name}")
            return element is not None and element.get(f"{{{W}}}val", "1") not in {"0", "false", "False"}

        reference = run.find(f"./{{{W}}}footnoteReference")
        fragment = {
            "текст": "".join(text.text or "" for text in run.findall(f".//{{{W}}}t")),
            "курсив": has_property("i"),
            "полужирный": has_property("b"),
            "сноска": None if reference is None else reference.get(f"{{{W}}}id"),
        }
        if fragment["сноска"] is not None and fragment["текст"]:
            fragments.extend([{**fragment, "сноска": None}, {**fragment, "текст": ""}])
        else:
            fragments.append(fragment)
    return style, fragments


def _footnotes(source: Path) -> dict[str, list[tuple[str, list[dict]]]]:
    with zipfile.ZipFile(source) as archive:
        try:
            root = ElementTree.fromstring(archive.read("word/footnotes.xml"))
        except KeyError:
            return {}
    footnotes = {}
    for footnote in root.findall(f"./{{{W}}}footnote"):
        identifier = footnote.get(f"{{{W}}}id")
        if identifier in {None, "-1", "0"}:
            continue
        footnotes[identifier] = [
            _footnote_fragments(paragraph)
            for paragraph in footnote.findall(f"./{{{W}}}p")
        ]
    return footnotes


def inspect_docx(path: Path) -> list[str]:
    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if len(entries) != len(names):
                return ["Документ DOCX повреждён или защищён паролем."]
            required_parts = {
                "[Content_Types].xml",
                "_rels/.rels",
                "word/document.xml",
                "word/_rels/document.xml.rels",
            }
            if not required_parts <= names:
                return ["Документ DOCX повреждён или защищён паролем."]
            xml_parts = [
                (entry.filename, ElementTree.fromstring(archive.read(entry)))
                for entry in entries
                if entry.filename.endswith((".xml", ".rels"))
            ]
            warnings = [message for name, message in UNSUPPORTED_PARTS.items() if name in names]
            found = set()
            contains_table = False
            for name, root in xml_parts:
                if not name.startswith("word/"):
                    continue
                for element in root.iter():
                    if element.tag == f"{{{W}}}tbl":
                        contains_table = True
                    local_name = element.tag.rsplit("}", 1)[-1]
                    if local_name in UNSUPPORTED_XML:
                        found.add(local_name)
            if contains_table:
                warnings.append("Документы с таблицами пока не поддерживаются.")
            return warnings + [
                message for name, message in UNSUPPORTED_XML.items() if name in found
            ]
    except (OSError, RuntimeError, zipfile.BadZipFile, ElementTree.ParseError):
        return ["Документ DOCX повреждён или защищён паролем."]


def extract_docx(source: Path, destination: Path) -> list[dict]:
    errors = inspect_docx(source)
    if errors:
        raise ValueError(" ".join(errors))
    document = Document(source)
    if document.tables:
        raise ValueError("Документы с таблицами пока не поддерживаются.")
    footnotes = _footnotes(source)
    inserted_footnotes = set()
    blocks = []
    for number, paragraph in enumerate(document.paragraphs, start=1):
        fragments = [fragment for run in paragraph.runs for fragment in _run_fragments(run)]
        blocks.append(
            {
                "идентификатор": f"B{number:06d}",
                "тип": _block_type(paragraph),
                "стиль": (paragraph.style.name or "") if paragraph.style else "",
                "фрагменты": fragments,
            }
        )
        for fragment in fragments:
            footnote_id = fragment["сноска"]
            if footnote_id not in footnotes or footnote_id in inserted_footnotes:
                continue
            inserted_footnotes.add(footnote_id)
            for index, (style, footnote_fragments) in enumerate(footnotes[footnote_id], start=1):
                blocks.append(
                    {
                        "идентификатор": f"F{footnote_id}-P{index}",
                        "тип": "сноска",
                        "стиль": style,
                        "фрагменты": footnote_fragments,
                    }
                )
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(
        json.dumps(blocks, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    return blocks


def load_blocks(path: Path) -> list[dict]:
    return json.loads(path.read_text(encoding="utf-8"))


def _replace_run_text(run, text: str) -> None:
    text_elements = list(run.iter(f"{{{W}}}t"))
    if not text_elements:
        if text:
            raise ValueError("Структура оформления документа изменилась.")
        return
    text_elements[0].text = text
    for element in text_elements[1:]:
        element.text = ""


def _rewrite_docx_package(path: Path, replacements: dict[str, bytes]) -> None:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as temporary:
        temporary_path = Path(temporary.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as target:
            for item in source.infolist():
                target.writestr(item, replacements.get(item.filename, source.read(item.filename)))
            for name, content in replacements.items():
                if name not in source.namelist():
                    target.writestr(name, content)
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


def _replace_footnote_text_in_package(path: Path, blocks: list[dict]) -> None:
    footnote_blocks = [block for block in blocks if block.get("тип") == "сноска"]
    if not footnote_blocks:
        return
    try:
        with zipfile.ZipFile(path) as archive:
            root = ElementTree.fromstring(archive.read("word/footnotes.xml"))
    except (KeyError, zipfile.BadZipFile, ElementTree.ParseError) as error:
        raise ValueError("Не удалось сохранить текст сносок документа.") from error

    for block in footnote_blocks:
        match = re.fullmatch(r"F(.+)-P(\d+)", str(block.get("идентификатор", "")))
        if match is None:
            raise ValueError("Структура сносок проверенного перевода некорректна.")
        footnote_id, paragraph_number = match.groups()
        footnote = root.find(f"./{{{W}}}footnote[@{{{W}}}id='{footnote_id}']")
        paragraphs = [] if footnote is None else footnote.findall(f"./{{{W}}}p")
        paragraph_index = int(paragraph_number) - 1
        if paragraph_index >= len(paragraphs):
            raise ValueError(f"Структура оформления блока {block['идентификатор']} изменилась.")
        fragments = _fragments(block)
        fragment_index = 0
        for run in paragraphs[paragraph_index].findall(f"./{{{W}}}r"):
            text = "".join(item.text or "" for item in run.iter(f"{{{W}}}t"))
            reference = _run_footnote_id(run)
            if text or reference is None:
                if fragment_index >= len(fragments):
                    raise ValueError(f"Структура оформления блока {block['идентификатор']} изменилась.")
                _replace_run_text(run, fragments[fragment_index]["текст"])
                fragment_index += 1
            if reference is not None:
                fragment_index += 1
        if fragment_index != len(fragments):
            raise ValueError(f"Структура оформления блока {block['идентификатор']} изменилась.")
    _rewrite_docx_package(
        path,
        {"word/footnotes.xml": ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)},
    )


def rebuild_docx(template: Path, translated_blocks: list[dict], destination: Path) -> None:
    if not isinstance(translated_blocks, list):
        raise ValueError("Проверенный перевод должен содержать список блоков.")
    with tempfile.TemporaryDirectory() as temporary_directory:
        expected_blocks = extract_docx(template, Path(temporary_directory) / "blocks.json")
    errors = validate_translation(expected_blocks, translated_blocks)
    if errors:
        raise ValueError(" ".join(errors))
    document = Document(template)
    main_blocks = [block for block in translated_blocks if block.get("тип") != "сноска"]
    if len(document.paragraphs) != len(main_blocks):
        raise ValueError("Количество блоков шаблона не совпадает с проверенным переводом.")
    for paragraph, block in zip(document.paragraphs, main_blocks, strict=True):
        if block.get("тип") == "разрыв_сцены":
            continue
        fragments = _fragments(block)
        fragment_index = 0
        for run in paragraph.runs:
            reference = _run_footnote_id(run)
            if run.text or reference is None:
                if fragment_index >= len(fragments):
                    raise ValueError(f"Структура оформления блока {block['идентификатор']} изменилась.")
                if reference is None:
                    run.text = fragments[fragment_index]["текст"]
                else:
                    _replace_run_text(run._r, fragments[fragment_index]["текст"])
                fragment_index += 1
            if reference is not None:
                fragment_index += 1
        if fragment_index != len(fragments):
            raise ValueError(f"Структура оформления блока {block['идентификатор']} изменилась.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    try:
        _replace_footnote_text_in_package(destination, translated_blocks)
        errors = inspect_docx(destination)
    except ValueError:
        destination.unlink(missing_ok=True)
        raise
    if errors:
        destination.unlink(missing_ok=True)
        raise ValueError(" ".join(errors))


def docx_has_footnote_reference(path: Path, footnote_id: str) -> bool:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))
    return any(
        reference.get(f"{{{W}}}id") == footnote_id
        for reference in root.iter(f"{{{W}}}footnoteReference")
    )


def docx_footnote_text(path: Path, footnote_id: str) -> str:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/footnotes.xml"))
    footnote = root.find(f"./{{{W}}}footnote[@{{{W}}}id='{footnote_id}']")
    if footnote is None:
        return ""
    return "".join(text.text or "" for text in footnote.iter(f"{{{W}}}t"))


def _fragments(block: dict) -> list[dict]:
    fragments = block.get("фрагменты", []) if isinstance(block, dict) else []
    return fragments if isinstance(fragments, list) else []


def _format_signature(block: dict) -> tuple:
    fragments = _fragments(block)
    return len(fragments), tuple(
        (
            fragment.get("курсив"),
            fragment.get("полужирный"),
            fragment.get("сноска"),
        )
        for fragment in fragments
        if isinstance(fragment, dict)
    )


def validate_translation(source_blocks: list[dict], translated_blocks: list[dict]) -> list[str]:
    errors = []
    expected = [block["идентификатор"] for block in source_blocks]
    actual = [block.get("идентификатор") if isinstance(block, dict) else None for block in translated_blocks]
    for identifier in expected:
        count = actual.count(identifier)
        if count == 0:
            errors.append(f"Блок {identifier} отсутствует.")
        elif count > 1:
            errors.append(f"Блок {identifier} повторён.")
    for identifier in actual:
        if identifier not in expected:
            errors.append(f"Обнаружен неизвестный блок {identifier}.")
    if actual != expected and len(actual) == len(expected) and all(identifier in expected for identifier in actual):
        errors.append("Порядок блоков изменён.")
    source_by_id = {block["идентификатор"]: block for block in source_blocks}
    for block in translated_blocks:
        if not isinstance(block, dict) or block.get("идентификатор") not in source_by_id:
            continue
        identifier = block["идентификатор"]
        if block.get("тип") != source_by_id[identifier].get("тип"):
            errors.append(f"Тип блока {identifier} изменён.")
        translated_fragments = block.get("фрагменты")
        if not isinstance(translated_fragments, list) or any(
            not isinstance(fragment, dict) or not isinstance(fragment.get("текст"), str)
            for fragment in translated_fragments
        ):
            errors.append(f"Фрагменты блока {identifier} содержат некорректный текст.")
            continue
        source_text = "".join(
            fragment.get("текст", "") for fragment in _fragments(source_by_id[identifier]) if isinstance(fragment, dict)
        )
        translated_text = "".join(
            fragment.get("текст", "") for fragment in _fragments(block) if isinstance(fragment, dict)
        )
        if source_text.strip() and not translated_text.strip():
            errors.append(f"Перевод блока {identifier} пуст.")
        if _format_signature(source_by_id[identifier]) != _format_signature(block):
            errors.append(f"Маркеры оформления блока {identifier} изменены.")
    return errors


def _block_length(block: dict) -> int:
    return sum(
        len(fragment.get("текст", ""))
        for fragment in _fragments(block)
        if isinstance(fragment, dict)
    )


def _referenced_footnotes(block: dict) -> set[str]:
    return {
        fragment["сноска"]
        for fragment in _fragments(block)
        if isinstance(fragment, dict) and isinstance(fragment.get("сноска"), str)
    }


def split_blocks(blocks: list[dict], max_chars: int) -> list[list[dict]]:
    if max_chars <= 0:
        raise ValueError("Максимальная длина части должна быть больше нуля.")
    groups = []
    index = 0
    while index < len(blocks):
        group = [blocks[index]]
        references = _referenced_footnotes(blocks[index])
        index += 1
        while index < len(blocks):
            candidate = blocks[index]
            identifier = candidate.get("идентификатор", "")
            if candidate.get("тип") != "сноска" or not any(identifier.startswith(f"F{reference}-P") for reference in references):
                break
            group.append(candidate)
            index += 1
        groups.append(group)

    chunks = []
    chunk = []
    size = 0
    for group in groups:
        group_size = sum(_block_length(block) for block in group)
        if chunk and size + group_size > max_chars:
            chunks.append(chunk)
            chunk = []
            size = 0
        chunk.extend(group)
        size += group_size
    if chunk:
        chunks.append(chunk)
    return chunks
