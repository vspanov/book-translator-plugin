import copy
import hashlib
import json
import os
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "skills" / "book-translator" / "scripts"
sys.path.insert(0, str(SCRIPTS))

import documents
import progress


def write_json(path: Path, value: dict | list) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def advance_successful_pipeline(
    project: Path,
    chapter: Path,
    source: list[dict],
    translated: list[dict],
    remaining: int,
    interrupt: bool = False,
    second_edit: bool = False,
) -> None:
    """Собирает все детерминированные артефакты без вызова модели."""
    chapter_work = project / "work" / chapter.stem
    chapter_work.mkdir(parents=True, exist_ok=True)
    source_path = chapter_work / "source.json"
    draft_path = chapter_work / "draft.json"
    edited_path = chapter_work / "edited-1.json"
    write_json(source_path, source)
    progress.advance_stage(project, "извлечение", str(source_path))
    write_json(draft_path, translated)
    progress.advance_stage(project, "перевод", str(draft_path))
    write_json(chapter_work / "completeness-1.json", {"ошибки": []})
    progress.advance_stage(project, "полнота_1", str(chapter_work / "completeness-1.json"))
    write_json(chapter_work / "report-1.json", {"статус": "пройдено", "замечания": []})
    progress.advance_stage(project, "проверка_1", str(chapter_work / "report-1.json"))
    write_json(edited_path, translated)
    progress.advance_stage(project, "редактура_1", str(edited_path))
    write_json(chapter_work / "completeness-2.json", {"ошибки": []})
    progress.advance_stage(project, "полнота_2", str(chapter_work / "completeness-2.json"))
    write_json(chapter_work / "report-2.json", {"статус": "пройдено", "замечания": []})
    progress.advance_stage(project, "проверка_2", str(chapter_work / "report-2.json"))
    if second_edit:
        progress.request_second_edit(project)
        edited_second_path = chapter_work / "edited-2.json"
        write_json(edited_second_path, translated)
        progress.advance_stage(project, "редактура_2", str(edited_second_path))
        write_json(chapter_work / "report-3.json", {"статус": "пройдено", "замечания": []})
        progress.advance_stage(project, "проверка_3", str(chapter_work / "report-3.json"))
    built = chapter_work / chapter.name
    documents.rebuild_docx(chapter, translated, built)
    progress.advance_stage(project, "сборка", str(built))
    next_state = chapter_work / "next-state"
    shutil.copytree(project / "state", next_state)
    with (next_state / "chapter-summaries.md").open("a", encoding="utf-8") as stream:
        stream.write(f"\n## {chapter.name}\nПринята тестовая глава.\n")
    progress.advance_stage(project, "память", str(next_state))
    progress.advance_stage(project, "фиксация", "подготовлена транзакция")
    next_progress = progress.load_progress(project)
    next_progress.update(
        {
            "статус_книги": "в_работе",
            "этап": "готово",
            "последняя_готовая_глава": chapter.name,
            "необработанных_глав": remaining,
            "ошибка": None,
        }
    )
    transaction = progress.prepare_transaction(project, chapter.name, built, next_state, next_progress)
    progress.commit_transaction(project, transaction, interrupt_after="state" if interrupt else None)


class PipelineIntegrationTests(unittest.TestCase):
    def test_three_chapters_publish_in_natural_order_recover_and_stop_cleanly(self):
        """Ловит потерю этапа, блока, фиксации или согласованности готовой книги."""
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory)
            input_dir = project / "input"
            input_dir.mkdir()
            input_hashes = {}
            for number in (1, 2, 10):
                document = Document()
                document.add_paragraph(f"Chapter {number}")
                chapter = input_dir / f"chapter-{number}.docx"
                document.save(chapter)
                input_hashes[chapter.name] = sha256(chapter)

            progress.initialize_project(project)
            chapters = documents.discover_chapters(project)
            self.assertEqual(["chapter-1.docx", "chapter-2.docx", "chapter-10.docx"], [chapter.name for chapter in chapters])
            manifest = documents.build_manifest(project, chapters)
            self.assertEqual(["chapter-1.docx", "chapter-2.docx", "chapter-10.docx"], [entry["имя"] for entry in manifest["главы"]])
            self.assertEqual(input_hashes, {entry["имя"]: entry["sha256"] for entry in manifest["главы"]})

            for index, chapter in enumerate(chapters):
                progress.start_chapter(project, chapter.name)
                source = documents.extract_docx(chapter, project / "work" / chapter.stem / "source.json")
                translated = copy.deepcopy(source)
                translated[0]["фрагменты"][0]["текст"] = f"Глава {index + 1}"
                self.assertEqual([], documents.validate_translation(source, translated))
                advance_successful_pipeline(
                    project,
                    chapter,
                    source,
                    translated,
                    remaining=len(chapters) - index - 1,
                    interrupt=index == 1,
                    second_edit=index == 0,
                )
                if index == 1:
                    self.assertFalse((project / "output" / chapter.name).exists())
                    draft_hash = sha256(project / "work" / chapter.stem / "draft.json")
                    progress.recover_transaction(project)
                    self.assertEqual(draft_hash, sha256(project / "work" / chapter.stem / "draft.json"))
                self.assertEqual([], progress.check_completed_chapter(project, chapter.name))
                published = documents.extract_docx(project / "output" / chapter.name, project / "work" / chapter.stem / "published.json")
                self.assertEqual([block["идентификатор"] for block in source], [block["идентификатор"] for block in published])
                self.assertEqual(f"Глава {index + 1}", published[0]["фрагменты"][0]["текст"])

            self.assertEqual([], documents.verify_manifest(project, manifest))
            self.assertEqual(input_hashes, {path.name: sha256(path) for path in input_dir.glob("*.docx")})
            progress.finish_book(project)
            self.assertEqual(
                ["chapter-1.docx", "chapter-2.docx", "chapter-10.docx"],
                [path.name for path in sorted((project / "output").glob("*.docx"), key=documents.natural_key)],
            )
            state = progress.load_progress(project)
            self.assertEqual("готово", state["статус_книги"])
            self.assertEqual("chapter-10.docx", state["последняя_готовая_глава"])
            self.assertEqual(0, state["необработанных_глав"])
            self.assertEqual([], progress.check_consistency(project))
            hook = subprocess.run(
                [sys.executable, str(ROOT / "hooks" / "check-progress.py")],
                input=json.dumps({"cwd": str(project)}),
                text=True,
                capture_output=True,
                check=False,
                env={**os.environ, "PYTHONDONTWRITEBYTECODE": "1"},
            )
            self.assertEqual(0, hook.returncode, hook.stderr)
            self.assertEqual({"continue": True, "suppressOutput": True}, json.loads(hook.stdout))

    def test_invalid_translation_keeps_published_output_and_memory_unchanged(self):
        """Ловит публикацию результата или памяти до успешной проверки полноты."""
        with tempfile.TemporaryDirectory() as directory:
            project = Path(directory)
            input_dir = project / "input"
            input_dir.mkdir()
            for number in (1, 2):
                document = Document()
                document.add_paragraph(f"Chapter {number}")
                document.save(input_dir / f"chapter-{number}.docx")
            progress.initialize_project(project)
            chapters = documents.discover_chapters(project)
            documents.build_manifest(project, chapters)
            first, second = chapters
            progress.start_chapter(project, first.name)
            source = documents.extract_docx(first, project / "work" / first.stem / "source.json")
            translated = copy.deepcopy(source)
            translated[0]["фрагменты"][0]["текст"] = "Глава 1"
            advance_successful_pipeline(project, first, source, translated, remaining=1)
            output_hash = progress.directory_sha256(project / "output")
            state_hash = progress.directory_sha256(project / "state")

            progress.start_chapter(project, second.name)
            source = documents.extract_docx(second, project / "work" / second.stem / "source.json")
            self.assertIn("отсутствует", " ".join(documents.validate_translation(source, [])))
            self.assertEqual(output_hash, progress.directory_sha256(project / "output"))
            self.assertEqual(state_hash, progress.directory_sha256(project / "state"))


if __name__ == "__main__":
    unittest.main()
