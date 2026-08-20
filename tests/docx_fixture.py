from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL = "http://schemas.openxmlformats.org/package/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"


def _rewrite_zip(path: Path, replacements: dict[str, bytes]) -> None:
    with NamedTemporaryFile(delete=False, suffix=".docx", dir=path.parent) as temporary:
        temporary_path = Path(temporary.name)
    with ZipFile(path, "r") as source, ZipFile(temporary_path, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            target.writestr(item, replacements.get(item.filename, source.read(item.filename)))
        for name, content in replacements.items():
            if name not in source.namelist():
                target.writestr(name, content)
    temporary_path.replace(path)


def make_formatted_docx(path: Path, *, footnote_in_text_run: bool = False) -> None:
    document = Document()
    document.add_heading("Chapter One", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Quiet ")
    italic = paragraph.add_run("thought")
    italic.italic = True
    bold = paragraph.add_run(" became an order.")
    bold.bold = True
    document.add_paragraph("* * *")
    document.add_paragraph("After the break.")
    document.save(path)

    with ZipFile(path, "r") as package:
        content_types = ET.fromstring(package.read("[Content_Types].xml"))
        ET.SubElement(content_types, f"{{{CT}}}Override", {
            "PartName": "/word/footnotes.xml",
            "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
        })
        relationships = ET.fromstring(package.read("word/_rels/document.xml.rels"))
        ET.SubElement(relationships, f"{{{REL}}}Relationship", {
            "Id": "rIdFootnotes",
            "Type": f"{R}/footnotes",
            "Target": "footnotes.xml",
        })
        body = ET.fromstring(package.read("word/document.xml"))
    paragraph_xml = body.findall(f".//{{{W}}}p")[-1]
    run = (
        paragraph_xml.findall(f"./{{{W}}}r")[-1]
        if footnote_in_text_run
        else ET.SubElement(paragraph_xml, f"{{{W}}}r")
    )
    ET.SubElement(run, f"{{{W}}}footnoteReference", {f"{{{W}}}id": "2"})
    footnotes = ET.fromstring(
        f'<w:footnotes xmlns:w="{W}"><w:footnote w:id="-1"/>'
        f'<w:footnote w:id="0"/><w:footnote w:id="2"><w:p><w:r>'
        f'<w:t>Original footnote.</w:t></w:r></w:p></w:footnote></w:footnotes>'
    )
    _rewrite_zip(path, {
        "[Content_Types].xml": ET.tostring(content_types, encoding="utf-8", xml_declaration=True),
        "word/_rels/document.xml.rels": ET.tostring(relationships, encoding="utf-8", xml_declaration=True),
        "word/document.xml": ET.tostring(body, encoding="utf-8", xml_declaration=True),
        "word/footnotes.xml": ET.tostring(footnotes, encoding="utf-8", xml_declaration=True),
    })
