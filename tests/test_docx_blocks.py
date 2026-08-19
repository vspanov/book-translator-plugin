import copy
import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from tests.docx_fixture import W as FIXTURE_W
from tests.docx_fixture import _rewrite_zip, make_formatted_docx


SCRIPTS = Path(__file__).resolve().parents[1] / "skills/book-translator/scripts"
sys.path.insert(0, str(SCRIPTS))
import documents


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/package/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"


def make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Chapter One", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Quiet ")
    italic = paragraph.add_run("thought")
    italic.italic = True
    bold = paragraph.add_run(" became an order.")
    bold.bold = True
    scene = document.add_paragraph("* * *")
    scene.style = document.styles["Normal"]
    document.add_paragraph("After the break.")
    document.save(path)


def add_footnote(path: Path) -> None:
    with zipfile.ZipFile(path) as archive:
        contents = {name: archive.read(name) for name in archive.namelist()}

    document = ElementTree.fromstring(contents["word/document.xml"])
    paragraphs = document.findall(f".//{{{W}}}p")
    reference_run = ElementTree.Element(f"{{{W}}}r")
    ElementTree.SubElement(reference_run, f"{{{W}}}footnoteReference", {f"{{{W}}}id": "7"})
    paragraphs[1].append(reference_run)
    contents["word/document.xml"] = ElementTree.tostring(document, encoding="utf-8", xml_declaration=True)

    relationships = ElementTree.fromstring(contents["word/_rels/document.xml.rels"])
    ElementTree.SubElement(
        relationships,
        f"{{{R}}}Relationship",
        {"Id": "rIdFootnotes", "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes", "Target": "footnotes.xml"},
    )
    contents["word/_rels/document.xml.rels"] = ElementTree.tostring(relationships, encoding="utf-8", xml_declaration=True)

    content_types = ElementTree.fromstring(contents["[Content_Types].xml"])
    ElementTree.SubElement(
        content_types,
        f"{{{CT}}}Override",
        {"PartName": "/word/footnotes.xml", "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"},
    )
    contents["[Content_Types].xml"] = ElementTree.tostring(content_types, encoding="utf-8", xml_declaration=True)

    footnotes = ElementTree.Element(f"{{{W}}}footnotes")
    for identifier in ("-1", "0"):
        ElementTree.SubElement(footnotes, f"{{{W}}}footnote", {f"{{{W}}}id": identifier})
    footnote = ElementTree.SubElement(footnotes, f"{{{W}}}footnote", {f"{{{W}}}id": "7"})
    paragraph = ElementTree.SubElement(footnote, f"{{{W}}}p")
    run = ElementTree.SubElement(paragraph, f"{{{W}}}r")
    ElementTree.SubElement(run, f"{{{W}}}t").text = "A note."
    contents["word/footnotes.xml"] = ElementTree.tostring(footnotes, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in contents.items():
            archive.writestr(name, content)


def block(identifier: str, text: str, *, footnote: str | None = None) -> dict:
    return {
        "идентификатор": identifier,
        "тип": "сноска" if identifier.startswith("F") else "абзац",
        "стиль": "Normal",
        "фрагменты": [{"текст": text, "курсив": False, "полужирный": False, "сноска": footnote}],
    }


class DocxBlockTests(unittest.TestCase):
    def test_rejects_document_with_table_before_writing_blocks(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "chapter.docx"
            target = root / "blocks.json"
            document = Document()
            document.add_paragraph("Before the table.")
            document.add_table(rows=1, cols=1).cell(0, 0).text = "Table text."
            document.save(source)

            with self.assertRaisesRegex(ValueError, "таблиц"):
                documents.extract_docx(source, target)

            self.assertFalse(target.exists())

    def test_extracts_stable_blocks_formatting_and_referenced_footnote(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "chapter.docx"
            target = root / "blocks.json"
            make_docx(source)
            add_footnote(source)

            blocks = documents.extract_docx(source, target)

            self.assertEqual(
                ["B000001", "B000002", "F7-P1", "B000003", "B000004"],
                [block["идентификатор"] for block in blocks],
            )
            self.assertEqual("заголовок", blocks[0]["тип"])
            self.assertTrue(blocks[1]["фрагменты"][1]["курсив"])
            self.assertTrue(blocks[1]["фрагменты"][2]["полужирный"])
            self.assertEqual("7", blocks[1]["фрагменты"][-1]["сноска"])
            self.assertEqual("сноска", blocks[2]["тип"])
            self.assertEqual("A note.", blocks[2]["фрагменты"][0]["текст"])
            self.assertEqual(blocks, documents.load_blocks(target))
            self.assertIn("заголовок", target.read_text(encoding="utf-8"))

    def test_rejects_missing_empty_duplicate_reordered_unknown_and_reformatted_blocks(self):
        source = [block("B000001", "One"), block("B000002", "Two")]
        variants = {
            "отсутствует": source[:1],
            "пуст": [{**source[0], "фрагменты": [{**source[0]["фрагменты"][0], "текст": ""}]}, source[1]],
            "повтор": [source[0], source[0], source[1]],
            "поряд": list(reversed(source)),
            "неизвест": source + [block("B999999", "Three")],
            "оформлен": [{**source[0], "фрагменты": [{**source[0]["фрагменты"][0], "курсив": True}]}, source[1]],
        }
        for word, translated in variants.items():
            with self.subTest(word=word):
                self.assertTrue(
                    any(word in error.lower() for error in documents.validate_translation(source, copy.deepcopy(translated)))
                )

    def test_rejects_changed_block_type(self):
        source = [{**block("B000001", "* * *"), "тип": "разрыв_сцены"}]
        translated = [block("B000001", "* * *")]

        errors = documents.validate_translation(source, translated)

        self.assertTrue(any("тип" in error.lower() for error in errors))

    def test_reports_null_fragment_text_without_crashing(self):
        source = [block("B000001", "One")]
        translated = [{**source[0], "фрагменты": [{**source[0]["фрагменты"][0], "текст": None}]}]

        errors = documents.validate_translation(source, translated)

        self.assertTrue(any("некоррект" in error.lower() for error in errors))

    def test_split_blocks_preserves_order_without_splitting_footnote_from_reference(self):
        blocks = [
            block("B000001", "x" * 20, footnote="7"),
            block("F7-P1", "x" * 20),
            block("B000002", "x" * 20),
            block("B000003", "x" * 20),
            block("B000004", "x" * 20),
        ]

        chunks = documents.split_blocks(blocks, max_chars=45)

        self.assertEqual([block["идентификатор"] for block in blocks], [block["идентификатор"] for chunk in chunks for block in chunk])
        self.assertEqual([2, 2, 1], [len(chunk) for chunk in chunks])
        self.assertEqual(["B000001", "F7-P1"], [block["идентификатор"] for block in chunks[0]])


class DocxRoundTripTests(unittest.TestCase):
    def test_rebuild_rejects_missing_or_reordered_footnote_block_before_writing(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            blocks_path = root / "blocks.json"
            make_formatted_docx(source)
            blocks = documents.extract_docx(source, blocks_path)
            variants = {
                "отсутствующая": [block for block in blocks if block["тип"] != "сноска"],
                "переставленная": [blocks[0], blocks[1], blocks[3], blocks[2], blocks[4]],
            }

            for name, translated in variants.items():
                with self.subTest(name=name):
                    result = root / f"{name}.docx"
                    with self.assertRaisesRegex(ValueError, "Блок|Порядок"):
                        documents.rebuild_docx(source, translated, result)
                    self.assertFalse(result.exists())

    def test_rebuild_preserves_supported_formatting_and_footnote_reference(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            result = root / "result.docx"
            blocks_path = root / "blocks.json"
            make_formatted_docx(source)
            blocks = documents.extract_docx(source, blocks_path)
            translated = copy.deepcopy(blocks)
            for block in translated:
                for fragment in block["фрагменты"]:
                    if fragment["текст"].strip():
                        fragment["текст"] = (
                            "Переведённая сноска"
                            if block["тип"] == "сноска" else "Перевод"
                        )
            self.assertEqual([], documents.validate_translation(blocks, translated))

            documents.rebuild_docx(source, translated, result)

            self.assertEqual([], documents.inspect_docx(result))
            rebuilt = Document(result)
            self.assertEqual("Heading 1", rebuilt.paragraphs[0].style.name)
            self.assertTrue(rebuilt.paragraphs[1].runs[1].italic)
            self.assertTrue(rebuilt.paragraphs[1].runs[2].bold)
            self.assertTrue(documents.docx_has_footnote_reference(result, "2"))
            self.assertIn("Переведённая сноска", documents.docx_footnote_text(result, "2"))

    def test_rebuild_preserves_reference_in_run_with_translated_text(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.docx"
            result = root / "result.docx"
            blocks_path = root / "blocks.json"
            make_formatted_docx(source, footnote_in_text_run=True)
            blocks = documents.extract_docx(source, blocks_path)
            final_block = next(block for block in blocks if block["идентификатор"] == "B000004")
            self.assertEqual(["After the break.", ""], [
                fragment["текст"] for fragment in final_block["фрагменты"]
            ])
            self.assertEqual([None, "2"], [
                fragment["сноска"] for fragment in final_block["фрагменты"]
            ])
            translated = copy.deepcopy(blocks)
            for block in translated:
                for fragment in block["фрагменты"]:
                    if fragment["текст"].strip():
                        fragment["текст"] = "Переведённая сноска" if block["тип"] == "сноска" else "Перевод"

            documents.rebuild_docx(source, translated, result)

            self.assertEqual("Перевод", Document(result).paragraphs[-1].runs[0].text)
            self.assertTrue(documents.docx_has_footnote_reference(result, "2"))


class DocxInspectionTests(unittest.TestCase):
    def test_inspect_rejects_incomplete_ooxml_package(self):
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "incomplete.docx"
            with zipfile.ZipFile(source, "w") as archive:
                archive.writestr("word/document.xml", f'<w:document xmlns:w="{W}"/>')

            self.assertEqual(
                ["Документ DOCX повреждён или защищён паролем."],
                documents.inspect_docx(source),
            )

    def test_extract_rejects_table_in_word_xml_part_before_writing_blocks(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "chapter.docx"
            target = root / "blocks.json"
            make_docx(source)
            _rewrite_zip(source, {
                "word/header1.xml": (
                    f'<w:hdr xmlns:w="{FIXTURE_W}"><w:tbl/></w:hdr>'.encode("utf-8")
                )
            })

            with self.assertRaisesRegex(ValueError, "таблиц"):
                documents.extract_docx(source, target)

            self.assertFalse(target.exists())


if __name__ == "__main__":
    unittest.main()
